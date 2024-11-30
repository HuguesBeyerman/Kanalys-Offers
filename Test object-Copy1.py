
import json
from datetime import datetime as dt
import click
from tabulate import tabulate
from pathlib import Path
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches,Pt,Length
from docx.oxml.ns import qn
from PyPDF2 import PdfWriter, PdfReader
from pdf2image import convert_from_path
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from PIL import Image
import subprocess
import PySimpleGUI as psg

class Offer:

    def __init__ (self):  
        self.offer_id = 'id'
        self.Customer = ''
        self.Building = ''
        self.Bom = {} 
        self.SOW = {}
        self.Path = 'offer_path' 
        self.Status = 'in approval'
        self.Created = dt.today().date().strftime('%d-%m-%y')
        self.Schema = ''

    def define_offer_id(self, offers_table):
        a = list(offers_table.keys())
        b = [i[:8] for i in a]
        c = b.count(dt.today().date().strftime('%d-%m-%y')) + 1
        self.offer_id = (dt.today().strftime('%d-%m-%y') + '-' + str(c))

    '''update offer table with new offer of updated offer'''
    def create_offer_table(self, offers_table, cust, build, bom_table, content, offer_path, path_schema):
        self.Customer = cust.cust_id
        self.Building = build.build_id
        self.Bom = bom_table
        self.SOW = content
        self.Path: offer_path 
        self.Schema: path_schema
        offers_table[self.offer_id] = {'Customer': self.Customer, 'Building': self.Building, 'Bom': self.Bom, 'SOW': self.SOW,
                                  'Path': self.Path, 'Status': self.Status,
                                  'Created': self.Created,
                                  'Schema': self.Schema}
        return(offers_table)
    
    def import_offer(self, id, offers_table):
        self.offer_id = id
        self.Customer = offers_table[id]['Customer']
        self.Building = offers_table[id]['Building']
        self.Bom = offers_table[id]['Bom']
        self.SOW = offers_table[id]['SOW']
        self.Path = offers_table[id]['Path']
        self.Status = offers_table[id]['Status']
        self.Created = offers_table[id]['Created']
        self.Schema = offers_table[id]['Schema']


class Customer:

    def append_cust_table(self, cust_table):
        cust_table[self.cust_id] = {'title':self.title, 'first_name': self.first_name,
                                        'last_name': self.last_name ,
                                        'company_name ': self.company_name,
                                        'email' : self.email,
                                        'phone_number':self.phone_number ,
                                        'street':self.street,
                                        'number':self.number,
                                        'postal_code':self.postal_code ,
                                        'city':self.city ,
                                        'country':self.country,
                                        'customer type':self.customer_type,
                                        'account_creation_date':self.account_creation_date,
                                        'VAT':self.VAT ,
                                        'Bank account':self.Bank_account}
        return(cust_table)

    def __init__ (self):
        self.cust_id = 'id'
        self.title = 'title'
        self.first_name = 'first_name'
        self.last_name = 'last_name'
        self.company_name = 'company_name'
        self.email = 'email'
        self.phone_number = 'phone_number'
        self.street = 'street'
        self.number = 'number'
        self.postal_code = 'postal_code'
        self.city = 'city'
        self.country = 'country'
        self.customer_type = 'customer type:'
        self.account_creation_date = dt.today().date().strftime('%d-%m-%y')
        self.VAT = 'VAT:'
        self.Bank_account = 'Bank account:'


    def define_cust_id(self,cust_table):
        id_list = list(cust_table.keys())
        id_list.sort(reverse = True, key = int)
        self.cust_id = str(int(id_list[0]) + 1)
    
    def encode_cust(self, title, first_name, last_name, company_name, email, phone_number, street, number, postal_code, city, country, customer_type, VAT, Bank_account):
        self.title = title
        self.first_name = first_name
        self.last_name = last_name
        self.company_name = company_name
        self.email = email
        self.phone_number = phone_number
        self.street = street
        self.number = number
        self.postal_code = postal_code
        self.city = city
        self.country = country
        self.customer_type = customer_type
        #self.account_creation_date = account_creation_date
        self.VAT = VAT
        self.Bank_account = Bank_account

    def import_cust(self, id, cust_table):
        self.cust_id = id
        try:
            self.title = cust_table[id]['title']
        except:
            self.title = 'Mr'
        self.first_name = cust_table[id]['first_name']
        self.last_name = cust_table[id]['last_name']
        self.company_name = cust_table[id]['company_name']
        self.email = cust_table[id]['email']
        self.phone_number = cust_table[id]['phone_number']
        self.street = cust_table[id]['street']
        self.number = cust_table[id]['number']
        self.postal_code = cust_table[id]['postal_code']
        self.city = cust_table[id]['city']
        self.country = cust_table[id]['country']
        try:
            self.customer_type = cust_table[id]['customer_type:']
        except:
            self.customer_type = 'prof'
        self.account_creation_date = cust_table[id]['account_creation_date']
        self.VAT = cust_table[id]['VAT']
        self.Bank_account = cust_table[id]['Bank account']

    def display_cust(self):
        b = [[i,j] for i,j in self.__dict__.items()]
        return(tabulate(b, headers=['Key', 'Value']))

    def __str__ (self):
        return ('{0}, {1}, {2}'.format(self.first_name, self.last_name, self.company_name))
    
    def export_cust(self):
        a = {self.cust_id : {
            'first_name': self.first_name, 
            'last_name': self.last_name, 
            'email': self.email, 
            'phone_number': self.phone_number , 
            'company_name': self.company_name,
            'customer_type': self.customer_type, 
            'account_creation_date': self.account_creation_date, 
            'VAT': self.VAT, 
            'Bank account': self.Bank_account, 
            'street': self.street, 
            'city': self.city, 
            'postal_code': self.postal_code, 
            'country': self.country, 
            'number': self.number, 
            'titre': self.title}}
        return(a)


class Building:    
    def __init__ (self):       
        self.build_id = ''
        self.street = 'Street'
        self.number = 'Number'
        self.postal_code = 'CP'
        self.city= 'City'
        self.country= 'Country'
        self.manager = 'cust_id'

    def encode_build (self, Street, Number, CP, City, Country, cust):       
        self.street = Street
        self.number = Number
        self.postal_code= CP
        self.city= City
        self.country= Country
        self.manager = cust.cust_id

    def define_build_id(self,build_table):
        try:
            id_list = list(build_table.keys())
            id_list.sort(reverse = True, key = int)
            self.build_id = str(int(id_list[0]) + 1)
        except:
            self.build_id = str(len(build_table.keys())+1)

    def encode_cust_address(self, cust):
        self.street = cust.street
        self.number =cust.number
        self.postal_code= cust.postal_code
        self.city= cust.city
        self.country= cust.country
        self.manager = cust.cust_id
 
    def import_build(self, id, build_table):
        self.build_id = id
        self.street = build_table[id]['street']
        self.number = build_table[id]['number']
        self.postal_code = build_table[id]['postal_code']
        self.city = build_table[id]['city']
        self.country = build_table[id]['country']
        self.manager = build_table[id]['manager']

    def __str__ (self):
        return ('{0}, {1}, {2}, {3}'.format(self.street, self.number, self.postal_code, self.city))

    def append_building_table(self, build_table):
        build_table[self.build_id] = {
                            'street': self.street,
                            'number' : self.number, 
                            'postal_code' : self.postal_code,
                            'city' : self.city,
                            'country' : self.country,
                            'manager' : self.manager}  
        return(build_table)

class Product:
    def __init__ (self):       
        self.prod_id = ''
        self.description = 'description'
        self.unit_selling_price = 'unit_selling_price'
        self.unit_cost = 'unit_cost'

    def define_prod_id(self, prod_table):
        id_list = list(prod_table.keys())
        id_list.sort(reverse = True, key = int)
        self.prod_id = str(int(id_list[0]) + 1)

    def encode_prod(self,description, unit_selling_price, unit_cost):
        self.description = description
        self.unit_selling_price = unit_selling_price
        self.unit_cost = unit_cost

    def import_prod(self, id, prod_table):
        self.prod_id = id
        self.description = prod_table[id]['description']
        self.unit_selling_price = prod_table[id]['unit_selling_price']
        self.unit_cost = prod_table[id]['unit_cost']

    def __str__ (self):
        return ('{0}, {1}, {2}'.format(self.description, self.unit_selling_price, self.unit_cost))

    def append_prod_table(self, prod_table):
            prod_table[self.prod_id] = {
                    'description': self.description,
                    'unit_selling_price' : self.unit_selling_price,
                    'unit_cost' : self.unit_cost}
            return(prod_table)

class Bom:
    def __init__ (self):
        self.bom_id = ''
        self.prod_id = 'product'
        self.description = 'description'
        self.unit_selling_price = 'unit_selling_price'
        self.qty = 'qty'
        self.discount = 'discount'
        self.line_price = 'line_price'

    def define_bom_id(self,bom_table):
        self.bom_id = len(bom_table) + 1

    def encode_bom(self, prod):
        self.prod_id = prod.prod_id
        self.description = prod.description
        self.unit_selling_price = prod.unit_selling_price
        
    def append_bom_table(self, bom_table):          
        bom_table[self.bom_id] = {
                    'prod_id': self.prod_id,
                    'description' : self.description,
                    'unit_selling_price' : self.unit_selling_price,
                    'qty' : self.qty,
                    'discount' : self.discount,                  
                    'line_price' : self.line_price}
        return(bom_table)


class Paragraph:
    def __init__ (self):
        self.parag_id = ''
        self.description = 'description'
        self.Paragraph = 'parag'
        self.elt = {}

    def define_parag_id(self, parag_table):
        self.parag_id = len(parag_table) + 1

    def import_parag(self, id, parag_table):
        self.parag_id = id
        self.description = parag_table[id]['description']
        self.Paragraph = parag_table[id]['Paragraph']
        self.elt = parag_table[id]['elt']

    def __str__ (self):
        return ('{0}'.format(self.description))

    def append_parag_table(self, parag_table):
        parag_table[self.parag_id] = {
            'description' : self.description,
            'Paragraph' : self.Paragraph,
            'elt' : self.elt}
        return(parag_table)

        
class Sow():
    def __init__ (self):
        self.parag_id = ''
        self.description = 'description'
        self.parag = 'parag'
        self.elt = {}

    def define_sow_id(self, content):
        self.sow_id = len(content) + 1

    def encode_sow(self, parag):
        self.parag_id = parag.parag_id
        self.description = parag.description
        self.parag = parag.Paragraph
        self.elt = parag.elt

    def append_sow(self, content):
        content[self.sow_id] = {
            'parag_id': self.parag_id,
            'description' : self.description,
            'Paragraph' : self.parag,
            'elt' : self.elt}
        return(content)
        

def open_table():
    with open("json/customers.json", 'r') as file:
        cust_table = json.load(file)
    #products and service    
    with open("json/sewer_products_services.json", 'r') as file:
        prod_table = json.load(file)
    # SOW paragraps
    with open("json/paragraphs.json", 'r') as file:
        parag_table = json.load(file)
    #buildings
    with open("json/buildings.json", 'r') as file:
        build_table = json.load(file)
    #offers    
    with open('json/offers.json', 'r') as file:
        offers_table = json.load(file)
        
    return(cust_table, build_table, prod_table, offers_table, parag_table)

def save_json(cust_table, build_table, offers_table, parag_table, prod_table):
    # Saving the list as a JSON file
    
    with open('json/customers.json', 'w') as file:
        json.dump(cust_table, file, indent=4)
    
    with open('json/offers.json', 'w') as file:
        json.dump(offers_table, file, indent=4)
        
    with open('json/buildings.json', 'w') as file:
        json.dump(build_table, file, indent=4)

    with open('json/sewer_products_services.json', 'w') as file:
        json.dump(prod_table, file, indent=4)

    with open('json/paragraphs.json', 'w') as file:
        json.dump(parag_table, file, indent=4)

    click.echo('Fichiers sauvés')
    
# Clear Screen
def clear_screen():
    input('\nPress Enter to continue.')
    if os.name == 'nt':  # For Windows
        os.system('cls')
    else:  # For macOS and Linux
        os.system('clear')

@click.command()
#Start
def start():
    cust_table, build_table, prod_table, offers_table, parag_table = open_table()
    
    x = True
    while x:
        click.clear()   
        choice = GUI.welcome_window()
        print(choice)
        
        if choice == '1':
            cust_table, build_table, prod_table, offers_table, parag_table = create_offer(cust_table, build_table, prod_table, offers_table, parag_table)

        elif choice == '2':
            #publish the offers selection   
            query = click.prompt('Search on customer name :', default='')
            if query == None:
                query = 'All'
            try:
                beg = min(dt.strptime(d['Created'], '%d-%m-%y') for d in offers_table.values())
                date_beg= click.prompt("Earliest date (format 'dd-mm-yy') ", default=beg)
                date_beg= dt.strptime(date_beg,'%d-%m-%y')
            except:
                date_beg = dt.strptime('01-01-00', '%d-%m-%y')

            try:
                lat = max(dt.strptime(d['Created'], '%d-%m-%y') for d in offers_table.values())
                date_end = click.prompt("Latest date (format 'dd-mm-yy') ", default=lat)
                date_end = dt.strptime(date_end, '%d-%m-%y')
            except:
                date_end = dt.today()
                
            offer_list = CLI.create_list_offers(offers_table, cust_table, s=query, date_beg=date_beg, date_end=date_end)
            
            #select offer to modify
            offer_id = click.prompt('\nWhich offer ID do you want to modify? ', type=click.Choice(offer_list))
            offer = Offer()
                
            if offers_table[offer_id]['Status'] != 'in approval': 
                offer.import_offer(offer_id, offers_table)
                offer.define_offer_id(offers_table)
                click.echo('This offer has already been sent. An new one will be created with ID nbr ' + str(offer.offer_id))
            else:
                offer.import_offer(offer_id, offers_table)
                
            cust_table, build_table, prod_table, offers_table, parag_table = modify_offer(offer, cust_table, build_table, prod_table, offers_table, parag_table)

        elif choice == '3':
            CLI.approve_offer(cust_table, build_table, prod_table, offers_table, parag_table)
            
        else:
            save_json(cust_table, build_table, offers_table, parag_table, prod_table)
            os.abort()
            x = False


'offer creation process'
def create_offer(cust_table, build_table, prod_table, offers_table, parag_table):
    offer = Offer()
    offer.define_offer_id(offers_table)
    cust_table, cust = CLI.select_client(cust_table, build_table, offers_table, parag_table, prod_table, offer)
    build_table, build = CLI.select_building(cust_table, build_table, offers_table, parag_table, prod_table, cust, offer)
    bom_table = {}
    prod_table, bom_table = CLI.select_bom(cust_table, build_table, offers_table, parag_table, prod_table, bom_table, offer)
    content = {}
    parag_table, content = CLI.select_content(cust_table, build_table, offers_table, parag_table, prod_table, content, offer)
    path_schema = crop_image(offer)
    z=True
    while z:
        
        bom_tabulate = tabulate_bom_table(bom_table)
        image = Image.open(path_schema)
        image.show()
        display_offre(bom_tabulate, content, offer, cust, parag_table, build)
        z, cust, cust_table, build, build_table, bom_table, content, prod_table, parag_table, path_schema = CLI.change_offer(cust, cust_table, build, build_table, bom_table, content, prod_table, parag_table, offer, path_schema, offers_table)

    bom_tabulate = tabulate_bom_table(bom_table)
    offer_path, path_schema = create_letter(cust_table, cust, bom_table, content, build_table, build, offer, bom_tabulate, parag_table, path_schema)
    offers_table = offer.create_offer_table(offers_table, cust, build, bom_table, content, offer_path, path_schema)
    return(cust_table, build_table, prod_table, offers_table, parag_table)


'offer modification process'
def modify_offer(offer, cust_table, build_table, prod_table, offers_table, parag_table):
    bom_table = offer.Bom
    print('import Baom_table: ', bom_table)
    content = offer.SOW
    cust = Customer()
    cust.import_cust(offer.Customer, cust_table)
    build = Building()
    build.import_build(offer.Building, build_table)
    path_schema = offer.Schema
    
    bom_tabulate = tabulate_bom_table(bom_table)
    image = Image.open(path_schema)
    image.show()
    display_offre(bom_tabulate, content, offer, cust, parag_table, build)
    z, cust, cust_table, build, build_table, bom_table, content, prod_table, parag_table, path_schema = CLI.change_offer(cust, cust_table, build, build_table, bom_table, content, prod_table, parag_table, offer, path_schema, offers_table)
    bom_tabulate = tabulate_bom_table(bom_table)
    offer_path, path_schema = create_letter(cust_table, cust, bom_table, content, build_table, build, offer, bom_tabulate, parag_table, path_schema)
    offers_table = offer.create_offer_table(offers_table, cust, build, bom_table, content, offer_path, path_schema)
    return (cust_table, build_table, prod_table, offers_table, parag_table)   

def start_test():
    open_table()
    cust = Customer()
    cust.define_cust_id(cust_table)
    cust_input = GUI.encode_cust(cust)
    cust.modify(title, first_name, last_name, company_name, email, phone_number, street, number, postal_code, city, country, customer_type, account_creation_date, VAT, Bank_account)
    click.echo('\nEncodé: ' + cust)

class GUI:

    def welcome_window():
        layout = [
            [psg.Button('Encode new offer', key='1'), 
             psg.Button('Modify existing offer', key='2'), 
             psg.Button('Approve offer', key='3'),
             psg.Button('Quit', key='-QUIT-')
            
        ]]

        window = psg.Window('Welcome', layout)       
        event, values= window.read()
        window.close()
        return event
        

    def customer_choice(offer):
        layout = [
            [psg.Button('Existing Customer', key='1'), 
             psg.Button('Encode new customer', key='2'), 
             psg.Button('Abort ', key='4')
            
        ]]
        
        window = psg.Window('Customer Choice - offer {}'.format(offer.offer_id), layout)       
        event, values= window.read()
        window.close()
        return event
    
    def encode_cust():
        layout = [
            [psg.Text('Title', size=(15,1)), psg.Input(key ='-TITLE-')],
            [psg.Text('First Name', size=(15,1)), psg.Input(key ='-FIRST_NAME-')],
            [psg.Text('Last Name', size=(15,1)), psg.Input(key ='-LAST_NAME-')],
            [psg.Text('Company Name', size=(15,1)), psg.Input(key ='-COMPANY_NAME-')],
            [psg.Text('Email', size=(15,1)), psg.Input(key ='-EMAIL-')],
            [psg.Text('Tel', size=(15,1)), psg.Input(key ='-TEL-')],
            [psg.Text('Street', size=(15,1)), psg.Input(key ='-STREET-')],
            [psg.Text('Number', size=(15,1)), psg.Input(key ='-NUMBER-')],
            [psg.Text('CP', size=(15,1)), psg.Input(key ='-CP-')],
            [psg.Text('City', size=(15,1)), psg.Input(key ='-CITY-')],
            [psg.Text('Country', size=(15,1)), psg.Input(key ='-COUNTRY-')],
            [psg.Text('Customer type', size=(15,1)), psg.Input(key ='-TYPE-')],
            [psg.Text('VAT', size=(15,1)), psg.Input(key ='-VAT-')],
            [psg.Text('BanK Account', size=(15,1)), psg.Input(key ='-BANK-')],
            [psg.Text('You entered ', size=(15,1)), psg.Text(key='-OUT-')],
            [psg.Button('OK', key='-OK-'), psg.Exit()],
        ]

        window = psg.Window('Customer', layout)
        while True:
            event, values = window.read()
            #print(event, values)
            out = '{} {} {}'.format(values['-FIRST_NAME-'], values['-LAST_NAME-'], values['-COMPANY_NAME-']) 
            window['-OUT-'].update(out)
            if event == psg.WIN_CLOSED or event == 'Exit':
                break
            elif event == '-OK-':
                False
                window.close()
                return(values)


    def select_cust(cust_table):
        # Initial customer list
        cust_list = [
            [i, j['first_name'], j['last_name'], j['company_name']]
            for i, j in cust_table.items()
        ]
    
        layout = [
            [psg.Text('Query:'), psg.Input(key='-QUERY-', enable_events=True)],
            [psg.Listbox(values=cust_list, enable_events=True, size=(40, 30), key="-CUST LIST-", select_mode=psg.LISTBOX_SELECT_MODE_SINGLE)]
        ]
    
        window = psg.Window('Customer Selector', layout)

        while True:
            event, values = window.read()
    
            if event == psg.WINDOW_CLOSED:
                break

            elif event == '-QUERY-':
                query = values['-QUERY-'].lower()  # Make query case-insensitive
                filtered_list = [
                    [i, j['first_name'], j['last_name'], j['company_name']]
                    for i, j in cust_table.items()
                    if query in (j.get('first_name', '').lower()) or 
                       query in (j.get('last_name', '').lower()) or 
                       query in (j.get('company_name', '').lower())
                ]
                # Update the Listbox with the filtered list
                window['-CUST LIST-'].update(filtered_list)
  
            # Handle selection from the Listbox
            elif event == '-CUST LIST-':
                selected_customer = values['-CUST LIST-'][0]  # Get selected item
                break 
            
        window.close()
        print(selected_customer[0])
        return selected_customer[0]  # Return the customer ID


    def cloture_session():
        layout = [
            [psg.Text('Do you confirm that you want to quit the prg? All previously encoded offers will be saved.')],
            [psg.Yes(), psg.No()]
        ]

        window = psg.Window('Abord session', layout)        
        event, values = window.read()
        window.close()
        return(event)

    def select_buidling(cust, build_table):
           
        build_list = [
            [i, j['street'], j['number'], j['postal_code'], j['city'], j['country']]
            for i, j in build_table.items() 
            if cust.cust_id == j['manager']         
        ]

        default_build = (build_list[0] if len(build_list) > 0 else ['', '', '', '', '', ''])

        layout = [
            [psg.Text('List of buildings of customer  {}'.format(str(cust)))],
            [psg.Listbox(values=build_list, enable_events=True, size=(55, 20), key="-BUILD LIST-", select_mode=psg.LISTBOX_SELECT_MODE_SINGLE)],
            [psg.Text('Street', size=(15,1)), psg.Input(key ='-STREET-', default_text=default_build[1])],
            [psg.Text('Number', size=(15,1)), psg.Input(key ='-NUMBER-', default_text=default_build[2])],
            [psg.Text('CP', size=(15,1)), psg.Input(key ='-CP-', default_text=default_build[3])],
            [psg.Text('City', size=(15,1)), psg.Input(key ='-CITY-', default_text=default_build[4])],
            [psg.Text('Country', size=(15,1)), psg.Input(key ='-COUNTRY-', default_text=default_build[5])],
            [psg.Button('OK', key='-OK-')]           
        ]
       
        window = psg.Window('Select Building', layout) 
        
        while True:
            event, values = window.read()
            if event == psg.WINDOW_CLOSED:
                break
            elif event == '-BUILD LIST-':
                print(values)
                default_build = values['-BUILD LIST-'][0]
                print(default_build)
                window['-STREET-'].update(default_build[1])
                window['-NUMBER-'].update(default_build[2])
                window['-CP-'].update(default_build[3])
                window['-CITY-'].update(default_build[4])
                window['-COUNTRY-'].update(default_build[5])
            elif event == '-OK-':
                False
                window.close()
                build = Building()
                build.encode_build(values['-STREET-'], values['-NUMBER-'], values['-CP-'], values['-CITY-'], values['-COUNTRY-'], cust)
                return(build)
            
                

class CLI:
    def encode_cust(cust):
        click.echo('Encode your customer\n-------------------\n')
        title = click.prompt('title (default: {}) '.format(cust.title), default=cust.title)
        first_name = click.prompt('first_name (default: {}) '.format(cust.first_name), default=cust.first_name)
        last_name = click.prompt('last_name (default: {}) '.format(cust.last_name), default=cust.last_name)
        company_name = click.prompt('company_name (default: {}) '.format(cust.company_name), default=cust.company_name)
        email = click.prompt('email (default: {}) '.format(cust.email), default=cust.email)
        phone_number = click.prompt('tel (default: {}) '.format(cust.phone_number), default=cust.phone_number)
        street = click.prompt('street (default: {}) '.format(cust.street), default=cust.street)
        number = click.prompt('number (default: {}) '.format(cust.number), default=cust.number)
        postal_code = click.prompt('postal_code (default: {}) '.format(cust.postal_code), default=cust.postal_code)
        city = click.prompt('city (default: {}) '.format(cust.city), default=cust.city)
        country = click.prompt('country (default: {}) '.format(cust.country), default=cust.country)
        customer_type = click.prompt('customer type: (default: {}) '.format(cust.customer_type), default=cust.customer_type, type=click.Choice(['Prof', 'Privé']))
        VAT = click.prompt('VAT: (default: {}) '.format(cust.VAT), default=cust.VAT)
        Bank_account = click.prompt('Bank account: (default: {}) '.format(cust.Bank_account), default=cust.Bank_account)
        return(title, first_name, last_name, company_name, email, phone_number, street, number, postal_code, city, country, customer_type, account_creation_date, VAT, Bank_account)

    def encode_build(build):
        click.echo('Encode your building\n-------------------\n')
        street = click.prompt('street (default: {}) '.format(build.street), default=build.street)
        number = click.prompt('number (default: {}) '.format(build.number), default=build.number)
        postal_code = click.prompt('CP (default: {}) '.format(build.postal_code), default=build.postal_code)
        city= click.prompt('city (default: {}) '.format(build.city), default=build.city)
        country = click.prompt('country (default: {}) '.format(build.country), default=build.country)
        return(street, number, postal_code, city, country)
            
    def select_client(cust_table, build_table, offers_table, parag_table, prod_table, offer):
        cust = Customer()
        choice = GUI.customer_choice(offer)   

        if choice == '1':
            #cust_id = click.prompt('\nSélectionner le Customer ID: ', type=click.Choice(list(cust_table.keys())))  
            cust_id = GUI.select_cust(cust_table)
            cust.import_cust(cust_id, cust_table)
            #click.echo('Result: ' + str(cust))
            #click.pause('Enter to continue')
            
        elif choice == '2':
            cust.define_cust_id(cust_table)
            cust_input = GUI.encode_cust()
            cust.encode_cust(cust_input['-TITLE-'],
                             cust_input['-FIRST_NAME-'],
                             cust_input['-LAST_NAME-'],
                             cust_input['-COMPANY_NAME-'],
                             cust_input['-EMAIL-'],
                             cust_input['-TEL-'],
                             cust_input['-STREET-'],
                             cust_input['-NUMBER-'],
                             cust_input['-CP-'],
                             cust_input['-CITY-'],
                             cust_input['-COUNTRY-'],
                             cust_input['-TYPE-'],
                             cust_input['-VAT-'],
                             cust_input['-BANK-'])
            
            cust_table = cust.append_cust_table(cust_table)
            build = Buidling()
            build.encode_build(cust.street, cust.number, cust.postal_code, cust.city, cust.country, cust)
            build_table = build.append_building_table(build_table)
            #click.echo('\nResult: ' + str(cust))
            #click.pause('Enter to continue')
            #i = False
                     
        else:
            cloture = GUI.cloture_session()
            if cloture == 'Yes':
                save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                os.abort()
            else:
                pass

        return(cust_table, cust)
        
    
    '''Select the building for the offer'''
    def select_building(cust_table, build_table, offers_table, parag_table, prod_table, cust, offer): 
        build = Building()
        build= GUI.select_buidling(cust, build_table)
        #print('build: ',build)
        #click.pause()
        #click.echo(tabulate([['Building - Offer ID: ' + str(offer.offer_id)]], tablefmt="simple_grid"))
    
        #build_qty, cust_address =  CLI.query_building(build_table, cust)
        x = True
        while x:
            exist = click.prompt('\n1.Existing Building 2.Encode new building 3.Encode customer Address 4.Abort ', type=click.Choice(['1','2','3','4']))
            if exist == '1' and len(build_qty) > 0:                    
                build_id = click.prompt('\nBuilding ID: ', type=click.Choice(build_qty))
                build.import_build(build_id, build_table)
                click.echo('Result: ' +  str(build))
                x = False
    
            elif exist == '1' and len(build_qty)== 0: 
                click.echo('\nNo Existing Building')
                    
            elif exist == '2':
                build.define_build_id(build_table)
                street, number, postal_code, city, country = CLI.encode_build(build)
                build.encode_build(street, number, postal_code, city, country, cust)
                build_table = build.append_building_table(build_table)
                x = False
    
            elif exist == '3' and len(cust_address)==0:
                build.define_build_id(build_table)
                build.encode_cust_address(cust)
                build_table = build.append_building_table(build_table)
                x = False
    
            elif exist == '3' and len(cust_address)>0:
                build.build_id = cust_address[0]
                build.encode_build(cust.street, cust.number, cust.postal_code, cust.city, cust.country, cust)
                x = False
    
            else:
                if click.confirm('Do you confirm that you want to quit the prg? All previously encoded offers will be saved.'):
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
    
        return(build_table, build)

    
    '''check if building is managed by client'''
    def match_build_cust(x, cust, build_table, build):
        if build_table[build.build_id]['manager'] != cust.cust_id:
            click.echo('The building is not managed by this client!')
        else:
            x = False
        return(x)

    ''' Takes a customer ID and searches the building data for buildings managed by that customer.
    Displays a list of buildings or indicates no buildings found.'''
    def query_building(build_table, cust):
        l = []
        #k = []
        #cust_address = []
        for i,j  in build_table.items():
            if cust.cust_id == j['manager']:
                m = [i]
                for x,y in j.items():
                    if x in ['street', 'number', 'postal_code', 'city']:
                        m.append(y)
                l.append(m)
                #k.append(i)
                #if m[1:] == [cust.street, cust.number, cust.postal_code, cust.city]:
                    #cust_address = m
        return(l)

    
    '''Provides options to add existing items, create new products, erase items, exit BOM creation or abort'''
    def select_bom(cust_table, build_table, offers_table, parag_table, prod_table, bom_table, offer):
        
        click.clear()
        click.echo(tabulate([['Bill of Materials- Offer ID :' + str(offer.offer_id)]], tablefmt="simple_grid"))
    
        x = True
        while x:
            choice = click.prompt('\n1: Add new item of pricelist, 2:Encode new product, 3:Erase item of the BOM, \n4:Search Product, 5: Exit, 6: Abort', type=click.Choice(['1','2','3','4','5','6']))

            if choice == '1':
                bom_table= CLI.select_prod_item(prod_table, bom_table)
                
            elif choice == '2':
                prod_table, bom_table = CLI.creation_prod_item(prod_table, bom_table)
                
            elif choice == '3':
                print(len(bom_table))
                if len(bom_table) == 0:
                    click.echo('\nNo item in the BOM')
                else:
                    #display bom items table
                    l=[]
                    k= []
                    for i, j in bom_table.items():
                        l.append([i, j['description'], j['unit_selling_price'], j['qty'], j['discount']])
                        k.append(i)
    
                    click.echo('Items in the offer Bill of Materials')
                    print('\n', tabulate(l, headers =['Bom ID', 'Description', 'Price', 'Qty', 'Disc' ]))
                    x= True
                    while x:
                        a = click.prompt('\nItem to erase: ', type= int)
                        if a in k:                        
                            del(bom_table[a])
                            click.echo('Item is erased.')
                            x = False
                 
            elif choice == '4':
                CLI.find_prod(prod_table)
                    
            elif choice == '5' and len(bom_table) == 0:
                click.echo('\nThere must be at least on item in the Bill of Material')
                
            elif choice == '6':
                if click.confirm('Do you confirm that you want to quit the prg? All previously encoded offers will be saved.'):
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                else:
                    continue
            
            else:
                x = False
    
                    
        return(prod_table, bom_table)   


        
    '''Allows adding existing products from a product list to the BOM.
    Takes user input for product ID and calculates the line price considering quantity and discount.'''
    def select_prod_item(prod_table, bom_table):
    
        a = click.prompt('\nProd ID: ', type=click.Choice(list(prod_table.keys())))
        prod = Product()
        bom = Bom()
        
        bom.define_bom_id(bom_table)
        prod.import_prod(a, prod_table)
        click.echo('Result : ' +  str(prod))
        bom.encode_bom(prod)
        bom.qty = click.prompt('Qty: ', type=int)
        bom.discount = click.prompt('Discount (%): ', type=click.IntRange(0, 100))
        bom.line_price = bom.qty * bom.unit_selling_price * (1 - bom.discount/100)          
        bom_table = bom.append_bom_table(bom_table)
        #print(bom_table)
        #print(len(bom_table))
        
        return(bom_table)


    '''Creates a new product entry with details like description, selling price, and cost price.
    Adds the new product to the BOM with specified quantity and discount.'''
    def creation_prod_item(prod_table, bom_table):
        
        prod = Product()
        bom = Bom()
        prod.define_prod_id(prod_table)
        prod.description = click.prompt('\nProd or Service Description: ')
        prod.unit_selling_price = click.prompt('Selling Price: ', type=float)
        prod.unit_cost = click.prompt('Cost: ', type=float)
        bom.define_bom_id(bom_table)
        bom.description = prod.description
        bom.unit_selling_price = prod.unit_selling_price
        bom.qty = click.prompt('Qty: ', type=int)
        bom.discount = click.prompt('Discount (%): ', type=click.IntRange(0, 100))               
        bom.line_price = bom.qty * bom.unit_selling_price * (1 - bom.discount/100)          
        bom_table = bom.append_bom_table(bom_table)
        prod_table = prod.append_prod_table(prod_table)
        #print(bom_table)
                
        return(prod_table, bom_table)


    ''' query product list: allows searching by description '''
    def find_prod(prod_table):
        click.echo('\nRecherche produit existant \n --------------------------')
        s = click.prompt('\nProduct query')
        l = []
        for i,j  in prod_table.items():
            if (s.lower() in (j.get('description') or '').lower()):
                l.append([i, j['description'], j['unit_selling_price'], j['unit_cost']])
        if len(l) > 0:
            print('\n', tabulate(l, headers =['Prod ID', 'Description', 'Price', 'Cost']))
        else:
            click.echo('\n no product')


    ''' Allows adding existing SOW paragraphs or creating new ones.
    Provides options to add paragraphs, create new paragraphs, erase paragraphs, or exit SOW creation.'''
    def select_content(cust_table, build_table, offers_table, parag_table, prod_table, content, offer):
    
        clear_screen()  
        click.echo(tabulate([['Statement of Work - Offer ID: ' +  str(offer.offer_id)]], tablefmt="simple_grid"))
    
        x = True
        while x:
            choice = click.prompt('\n1:Add new SOW item , 2:create and add a new SOW item, 3:Erase item of the SOW,\n 4:Search SOW item, 5:Exit, 6: Abort\n', type=click.Choice(['1','2','3','4','5','6']))
            if choice == '1':
                content = CLI.select_sow(parag_table, content)
                
            elif choice == '2':
                parag_table, content= CLI.creation_sow_item(parag_table, content)
                
                
            elif choice == '3':
                if len(content) == 0:
                    click.echo('\nNo item in the SOW')

                else:
                    #display bom items table
                    l=[]
                    k= []
                    for i, j in content.items():
                        l.append([i, j['description'], j['Paragraph'][:40]])
                        k.append(i)
    
                    click.echo('List of paragraphs of the offer\n--------------------------------')
                    print('\n', tabulate(l, headers =['Content ID', 'Description', 'Paragraph']))
                    x= True
                    while x:
                        a = click.prompt('\nItem to erase: ', type= int)
                        if a in k:                        
                            del(content[a])
                            click.echo('Item is erased.')
                            x = False                

            elif choice == '4':
                CLI.find_parag(parag_table)
                    
            elif choice == '5' and len(content) == 0:
                click.echo('\nThere must be at least one item in the statement of work')

            elif choice == '6':
                if click.confirm('Do you confirm that you want to quit the prg? All previously encoded offers will be saved.'):
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                else:
                    continue
                
            else:
                x = False
        
        return(parag_table, content)

    
    '''Select SOW item of the list to be added to offer'''
    def select_sow(parag_table, content):
        
        parag = Paragraph()
        sow = Sow()        
        a = click.prompt('\nParagraph ID: ', type=click.Choice(list(parag_table.keys())))

        sow.define_sow_id(content)
        parag.import_parag(a, parag_table)
        click.echo('Result : ' +  str(parag))
        sow.encode_sow(parag)
        for i,j in sow.elt.items():
            sow.elt[i] = click.prompt(' Encode {0}: '.format(i), default = j)
        content = sow.append_sow(content)
          
        return(content)

    
    '''Creates a new SOW paragraph entry with a description.'''
    def creation_sow_item(parag_table, content):
    
        parag = Paragraph()
        sow = Sow() 
        parag.define_parag_id(parag_table)
        sow.define_sow_id(content)
        parag.description = click.prompt('\nSOW Description: ')
        sow.description = parag.description
        parag.Paragraph = click.prompt('\nSOW Paragraph: ')
        sow.Paragraph = parag.Paragraph
        elt_count = parag.Paragraph.count('{}')
        print('nbre elt: ', str(elt_count))
        for i in range(elt_count):
            clef = click.prompt('Element title ' + str(i+1) + ': ')
            parag.elt[clef] = ''
            sow.elt[clef] = click.prompt('Encode value of ' + clef + ': ')
        parag_table = parag.append_parag_table(parag_table)
        content = sow.append_sow(content)        
      
        return(parag_table, content)


    ''' query sow list: allows searching by description and paragraph content '''
    def find_parag(parag_table):
        click.echo('\nRecherche paragraph existant')
        click.echo('----------------------------')
        s = click.prompt('Paragraph query:')
        l = []
        for i,j  in parag_table.items():
            if (s.lower() in (defa).lower() or s.lower() in (j.get('Paragraph') or '').lower()):
                l.append([i, j['description'], j['Paragraph'][:40]])
        if len(l) > 0:
            print('\n', tabulate(l, headers =['Parag ID', 'Description', 'Paragraph']))
        else:
            click.echo('\n no content')


    ''' function to change offer content (customer, building, bom items, sow content'''
    def change_offer(cust, cust_table, build, build_table, bom_table, content, prod_table, parag_table, offer, path_schema, offers_table):
    
        x = True
        while x:    
            choice = click.prompt('\nDo you want to change the content of this offer?\n1(Customer), 2(Building), 3(prod & service), 4(job description),\n 5(schema), 6(display update), 7(create letter and exit)\n', type=click.Choice(['1', '2', '3', '4', '5', '6', '7']))
    
            if choice == '1':
                click.echo('\nYour choice: ' + str(cust))
                cust_table, cust = CLI.select_client(cust_table, build_table, offers_table, parag_table, prod_table, offer)
    
                y = True
                while y:
                    y = CLI.match_build_cust(y, cust, build_table, build)
                    if y:
                        #build_qty, cust_address = CLI.query_building(build_table, cust)
                        build_table, build = CLI.select_building(cust_table, build_table, offers_table, parag_table, prod_table, cust, offer)
                        y = False
    
            elif choice == '2':
                click.echo('\nYour choice: ' + str(build))
                #build_qty, cust_address = CLI.query_building(build_table, cust)
                build_table, build = CLI.select_building(cust_table, build_table, offers_table, parag_table, prod_table, cust, offer)
    
            elif choice == '3':
                header = list(list(bom_table.values())[0].keys())
                header[1] = 'Unit Price'
                bom_list= []
                for i, j in bom_table.items():
                    a = list(str(i)) + list(j.values())[1:4]
                    bom_list.append(a)
    
                print('\nYour choice: \n', tabulate(bom_list, headers = header))
                prod_table, bom_table = CLI.select_bom(cust_table, build_table, offers_table, parag_table, prod_table, bom_table, offer)
                
            elif choice == '4':
                a = [[i, j['description']] for i,j in content.items()]
                print('\nYour choice: \n', tabulate(a, headers = ['SOW ID', 'Description']))
                parag_table, content = CLI.select_content(cust_table, build_table, offers_table, parag_table, prod_table, content, offer)

            elif choice == '5':
                path_schema = crop_image(offer)
                image = Image.open(path_schema)
                image.show()
                
            elif choice == '6':
                bom_tabulate = tabulate_bom_table(bom_table)
                display_offre(bom_tabulate, content, offer, cust, parag_table, build)
    
            elif choice == '7':       
                x = False
                z = False
    
    
        return(z, cust, cust_table, build, build_table, bom_table, content, prod_table, parag_table, path_schema) 


    ''' list of offers'''
    def create_list_offers(offers_table, cust_table, status = 'All', s= 'All', date_beg=dt.strptime('01-01-20', '%d-%m-%y'),
                           date_end=dt.today()):
    
        cu = []
        for i,j  in cust_table.items():
            if (s.lower() in (j.get('last_name') or '').lower())  or (s.lower()  in (j.get('company_name') or '').lower())  or (s.lower()  in (j.get('first_name') or '').lower()):
                cu.append(i)
        t = []
        offer_list = []
        if len(cu) > 0 or s == 'All':
              
            for i, j in offers_table.items():
                if ((status == 'All'or j['Status'] == status )
                    and (s == 'All' or j['Customer'] in cu)
                    and (dt.strptime(j['Created'], '%d-%m-%y') >= date_beg)
                    and (dt.strptime(j['Created'], '%d-%m-%y') <= date_end)):
                                        
                                                                                      
                    l =[i]
                    offer_list.append(i)
                    if cust_table[j['Customer']]['company_name'] == ' ':
                        l.append(cust_table[j['Customer']]['last_name'])
                    else:
                        l.append(cust_table[j['Customer']]['company_name'])
    
                    c = 0    
                    for a,b in j['Bom'].items():
                        c +=  b['line_price']
                    l.append(c)
                    l.append(j['Status'])
                    l.append(j['Created'])
                    t.append(l)
        else:
            click.echo('No identified customers.')
            
        click.echo('\nOffer list :')    
        print(tabulate(t, headers = ['Id', 'Client', 'Value', 'Status', 'last update']))
        return(offer_list)


    'approve/modify offers'
    def approve_offer(cust_table, build_table, prod_table, offers_table, parag_table):
        clear_screen()
        click.echo(tabulate([['Offer Approval: ']], tablefmt="simple_grid"))
        CLI.create_list_offers(offers_table, cust_table, status = 'in approval')
    
    
        x = True
        while x:
            choice = click.prompt('\n1: Display Offer, 2: Modify offer, 3:Erase offer ,\n 4:Approve offer, 5:Exit, 6: Abort\n', type=click.Choice(['1','2','3','4','5','6']))
            offer = Offer()
            
            if choice == '1':
                offer_id = click.prompt('\nOffer ID: ', type=click.Choice(list(offers_table.keys())))
                offer.import_offer(offer_id, offers_table)
                bom_tabulate = tabulate_bom_table(offer.Bom)

                cust = Customer()
                cust.import_cust(offer.Customer, cust_table)
                build = Building()
                build.import_build(offer.Building, build_table)

                display_offre(bom_tabulate, offer.SOW, offer, cust, parag_table, build)
                image = Image.open(offer.Schema)
                image.show()
                
            elif choice == '2':
                offer_id = click.prompt('\nOffer ID: ', type=click.Choice(list(offers_table.keys())))
                offer.import_offer(offer_id, offers_table)
                if offer.Status != 'in approval':
                    offer.define_offer_id(offers_table)
                    click.echo('This offer has already been sent. An new one will be created with ID nbr ' + str(offer.offer_id))
                    modify_offer(offer, cust_table, build_table, prod_table, offers_table, parag_table)
                    
                else:
                    modify_offer(offer, cust_table, build_table, prod_table, offers_table, parag_table)
            
                
            elif choice == '3':
                offer_id = click.prompt('\nOffer to erase: ', type=click.Choice(list(offers_table.keys())))
                if offers_table[offer_id]['Status'] == 'in approval':
                    del(offers_table[offer_id])
                    click.echo('Item erased.')
                else:
                    click.echo('Already ben sent. May not be deleted')

            elif choice == '4':
                offer_id = click.prompt('\nOffer to approve: ', type=click.Choice(list(offers_table.keys())))
                if offers_table[offer_id]['Status'] == 'in approval':
                    offers_table[offer_id]['Status'] = 'approved'
                    click.echo('Approved')
                    #collect data for mail
                    offer.import_offer(offer_id, offers_table)
                    cust = Customer()
                    cust.import_cust(offer.Customer, cust_table)
                    body = cust.title + ' ' + cust.first_name + ' ' + cust.last_name + '\n' + parag_table['17']['Paragraph'].format(cust.street, cust.number, cust.postal_code, cust.city)
                    print(offer.Path, type(offer.Path))
                    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', offer.Path, '--outdir', 'Letters'])
                    attach = offer.Path[:-4] + 'pdf'
                    print('attachement: ', attach)
                    titre = 'Kanalis offer ' + str(offer.offer_id)
                    send_offer(cust.title, cust.email, body, attach)
                    click.echo('Offer '+ str(offer.offer_id) + ' is sent!')
                else:
                    click.echo('Offer not in approval status')
                    
            elif choice == '6':
                if click.confirm('Do you confirm that you want to quit the prg? All previously encoded offers will be saved.'):
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                else:
                    continue
                
            else:
                x = False
    
    
        return(cust_table, build_table, offers_table, parag_table, prod_table)  
            

'''crop image from visit document'''
def crop_image(offer):
    
    clear_screen()
    click.echo(tabulate([['Add Schema- Offer ID: ' +  str(offer.offer_id)]], tablefmt="simple_grid"))
    
    #identify available reports and schema
    click.echo('\nlist of availabe reports and snapshots\n---------------------------------------')
    os.chdir('Chantiers')
    click.echo(os.listdir())

    y = True
    while y:
        path_schema = input('\nEncode file to load: ')
        if path_schema not in os.listdir():
            click.echo('This file does not exist')
        else:
            y = False
    
    type_image = click.prompt('\n1.image (jpg, png) , 2.rapport standard (pdf) ', type=click.Choice(['1', '2']))

    x = True
    while x: 
        if type_image == '1':
            if path_schema[-3:] not in ['jpg', 'png']:
                click.echo(' It is not the goof format')
            else:
                x = False

        #crop image from pdf report
        elif type_image == '2':
            if path_schema[-3:] != 'pdf':
                click.echo(' It is not the goof format')
                
            else:
                reader = PdfReader(path_schema)
                page = reader.pages[0]

                writer = PdfWriter()

                page.cropbox.upper_left = ( 100, 220)
                page.cropbox.lower_right = ( 350, 50)
                writer.add_page(page) 
          
                with open('result.pdf','wb') as fp:
                    writer.write(fp)


                image = convert_from_path('result.pdf', use_cropbox=True)
                image[0].save('Image.jpg', 'JPEG')
                path_schema = 'Image.jpg'
                os.remove("result.pdf")
                x= False

    #load final schema in dedicated schma folder and return path
    new_path_schema = 'Schema/schema '+ str(offer.offer_id) + path_schema[-4:]
    os.chdir('..')
    os.rename('Chantiers/{0}'.format(path_schema), new_path_schema)
    return(new_path_schema)


'''create BOM table with pricing that will be integrated in the offer display and letter'''
def tabulate_bom_table(bom_table):
    header = list(list(bom_table.values())[0].keys())[1:]
    header[1] = 'Unit Price'
    header[3] = 'Disc'
    bom_list= []
    for i in bom_table.values():
        a = list(i.values())[1:]
        bom_list.append(a)
    #bom_list
    
    for i in bom_list:
        i[1] = "{:.2f} €".format(i[1])
        i[3] = "{:.0f} %".format(i[3])
        i[4] = "{:.2f} €".format(i[4])
    
    # calcul total
    tot_htva = 0
    tva = 0
    for j in bom_table.values():
        tot_htva +=  j['line_price']
    tva += tot_htva * 0.21 
    tot_tvac = tot_htva * 1.21
    bom_list.append(['-----', '-----', '---', '-----', '-----'])
    bom_list.append([' ', ' ', ' ', 'Total HTVA', "{:.2f} €".format(tot_htva)])
    bom_list.append([' ', ' ', ' ', 'TVA', "{:.2f} €".format(tva)])
    bom_list.append([' ', ' ', ' ', 'Total TVAC', "{:.2f} €".format(tot_tvac)])
    bom_tabulate = tabulate(bom_list, headers = header, tablefmt="orgtbl", colalign = ('left', 'right', 'left' , 'left' ,'right'), floatfmt="{:.2f} €", maxcolwidths=[15, 10, 6, 12, 10])            
    return(bom_tabulate)


''' display offer content on screen'''
def display_offre(bom_tabulate, content, offer, cust, parag_table, build):

    clear_screen()
    click.echo(tabulate([['Offer content']], tablefmt="simple_grid"))
    cust_ad = '\t\t\t\t\t\t {0} \n\t\t\t\t\t\t {1} {2} \n\t\t\t\t\t\t {3} {4} \n\t\t\t\t\t\t {5} {6} \n\t\t\t\t\t\t {7}'.format(cust.company_name, cust.first_name, cust.last_name, cust.street, cust.number, cust.postal_code, cust.city, cust.country)
    click.echo(cust_ad + '\n')
    click.echo("Reference: " + str(offer.offer_id) + '\n')
    click.echo('\t\t\t\t\t\t' + dt.today().strftime('%B %d, %Y') + '\n')
    click.echo("Dear " + cust.title + ' ' + cust.first_name + ' ' + cust.last_name + ',\n')
    click.echo(parag_table['10']['Paragraph'].format(build.street, build.number, build.postal_code, build.city) + '\n')
    click.echo('Description des travaux\n------------------\n')
    for j in content.values():
        click.echo(j['Paragraph'].format(*[i for i in list(j['elt'].values())]) + '\n')
    click.echo('Bill of Materials\n------------------\n')
    click.echo (bom_tabulate + '\n')
    click.echo(parag_table['11']['Paragraph'] + '\n')
    click.echo("Sincerely,\n\nYour Name\nYour Job Title")
    

def send_offer(titre, email, body, attach):   
    # creates SMTP session
    s = smtplib.SMTP('smtp.gmail.com', 587)
    # start TLS for security
    s.starttls()
    # Authentication
    s.login("beyerman@gmail.com", "invf fzrx lgwn pncd")

    # Create a MIMEText object to represent the email
    msg = MIMEMultipart()
    msg['From'] = 'beyerman@gmail.com'
    msg['To'] = email
    msg['Subject'] = titre
    msg.attach(MIMEText(body, 'plain'))

    with open(attach,'rb') as file:
        # Attach the file with filename to the email
        msg.attach(MIMEApplication(file.read(), Name= titre + '.pdf'))
        
    # sending the mail
    s.sendmail("beyerman@gmail.com", email, msg.as_string())
    # terminating the session
    s.quit()


#write offer to docx document
def create_letter(cust_table, cust, bom_table, content, build_table, build, offer, bom_tabulate, parag_table, path_schema):

    # Create a new Document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier'
    font.size = Pt(10)
    
    # Add logo in the left cell
    logo = doc.add_paragraph()
    logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pict_logo = logo.add_run()
    pict_logo.add_picture('Kanalis-logo-ok.png', width=Inches(1))
    
    # Add sender's address in the right cell
    sender_paragraph = doc.add_paragraph("Kanalys SPRL\nRue de la Croix 41 \n1050 Ixelles", style ='Body Text')
    sender_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add a line break
    doc.add_paragraph()
    
    # Add recipient's address on the right below
    cust_ad = '\t\t\t\t\t\t {0} \n\t\t\t\t\t\t {1} {2} \n\t\t\t\t\t\t {3} {4} \n\t\t\t\t\t\t {5} {6} \n\t\t\t\t\t\t {7}'.format(cust.company_name, cust.first_name, cust.last_name, cust.street, cust.number, cust.postal_code, cust.city, cust.country)
    recipient_paragraph = doc.add_paragraph(cust_ad)
    recipient_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #recipient_address = recipient_paragraph.add_run(cust_ad)
    
    # Add today's date on the right below the recipient's address
    doc.add_paragraph()  # Add some space
    date_paragraph = doc.add_paragraph('Brussels, ' + dt.today().strftime('%B %d, %Y'))
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add reference section below the date
    doc.add_paragraph()  # Add some space
    reference_paragraph = doc.add_paragraph()
    reference_text = reference_paragraph.add_run("Reference: " + str(offer.offer_id))
    
    # Add greeting and the main letter content
    doc.add_paragraph()  # Add some space
    greeting_paragraph = doc.add_paragraph("Dear " + cust.title + ' ' + cust.first_name + ' ' + cust.last_name + ',')
    greeting_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Main text of the letter
    doc.add_paragraph(
        parag_table['10']['Paragraph'].format(build.street, build.number, build.postal_code, build.city))


    # Add work description
    # Add a title or heading
    doc.add_heading('Description des travaux', level=1)

    ligne = 30
    
    for j in content.values():
            ligne += round(len(j['Paragraph']))
            paragraph = doc.add_paragraph(j['Paragraph'].format(*[i for i in list(j['elt'].values())]))
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_together

    #Add picture
    try:
        ligne += 17
        doc.add_heading('Schema des travaux', level=1)
        doc.add_picture(path_schema, height=Inches(2))
    except UnrecognizedImageError:
        click.echo('Image non compatible.')
    
    # Add a table with two columns (Index, Paragraph)
    doc.add_heading('Bill of Materials', level=1)
        
    # Add Bom table
    paragraph = doc.add_paragraph(bom_tabulate)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.keep_together
        
    # Add closing remarks
    closing_paragraph = doc.add_paragraph()
    closing_text = closing_paragraph.add_run('\n' + parag_table['11']['Paragraph'] + "\nSincerely,\n\nYour Name\nYour Job Title")
    #closing_text.font.size = Pt(12)
    
    # Save the document
    name_offer = offer.offer_id + '.docx'
    os.chdir('Letters')
    doc.save(name_offer)
    offer_path = 'Letters/' + name_offer
    click.echo("\nLetter template created successfully.")
    os.chdir('..')
    return(offer_path, path_schema)

if __name__ == "__main__":
   start()


