#!/usr/bin/env python
# coding: utf-8

# In[63]:


#open modules
import json
from datetime import datetime as dt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches,Pt,Length
from docx.oxml.ns import qn
#from docx.oxml import OxmlElement
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from tabulate import tabulate
from pathlib import Path
import os
from PyPDF2 import PdfWriter, PdfReader
from pdf2image import convert_from_path
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from PIL import Image
#from docx2pdf import convert
import subprocess

# In[64]:


# Open and load the JSON file
#prospects
with open("json/customers.json", 'r') as file:
    cust_table = json.load(file)
#products and service    
with open("json/sewer_products_services.json", 'r') as file:
    prod_table = json.load(file)
# SOW paragraps
with open("json/paragraphs.json", 'r') as file:
    parag_table = json.load(file)
#buildings
with open("json/buildings.json", 'r') as file:
    build_table = json.load(file)
#offers    
with open('json/offers.json', 'r') as file:
    offers_table = json.load(file)


# Clear Screen
def clear_screen():
    input('\nPress Enter to continue.')
    if os.name == 'nt':  # For Windows
        os.system('cls')
    else:  # For macOS and Linux
        os.system('clear')


#Start
def start(cust_table, build_table, prod_table, offers_table, parag_table):

    x = True
    while x:
        clear_screen()    
        print(tabulate([['Kanalys Offers']], tablefmt="heavy_grid"))
        choice = input('1.Encode new offer, 2.Modify existing offer, 3.Approve offer, 4.Quit ')
        
        if choice not in ['1', '2', '3', '4']:
            print('Encode the right answer.')
        else:
            if choice == '1':
                cust_table, build_table, prod_table, offers_table, parag_table = create_offer(cust_table, build_table, prod_table, offers_table, parag_table)
 
            elif choice == '2':
                #publish the offers selection   
                query = input('Search on customer name :')
                if query == None:
                    query = 'All'
                try:    
                    date_beg= dt.strptime(input("Earliest date (format 'dd-mm-yy') "),'%d-%m-%y')
                except ValueError:
                    date_beg = dt.strptime('01-01-00', '%d-%m-%y')

                try:
                    date_end = dt.strptime(input("Latest date (format 'dd-mm-yy') "), '%d-%m-%y')
                except ValueError:
                    date_end = dt.today()
                    
                offer_list = create_list_offers(offers_table, cust_table, s=query, date_beg=date_beg, date_end=date_end)
                
                #select offer to modify
                offer_id = input('\nWhich offer ID do you want to modify? ')
                y = True

                #check if offer_id exist and create new offer if not in approval  
                if offer_id in offer_list:
                    
                    if offers_table[offer_id]['Status'] != 'in approval':
                        new_offer_id = create_offer_nbr(offers_table)
                        print('This offer has already been sent. An new one will be created with ID nbr ' + new_offer_id)
                        offers_table['new_offer_id'] = offers_table['new_offer_id']
                    else:
                        new_offer_id = offer_id
                        
                    cust_table, build_table, prod_table, offers_table, parag_table = modify_offer(new_offer_id, cust_table, build_table, prod_table, offers_table, parag_table)

                else:
                    print('This ID is not in the Offer list')
                    y = False

            elif choice == '3':
                approve_offer(cust_table, build_table, prod_table, offers_table, parag_table)
                
            else:
                save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                os.abort()
                x = False




"""Attributes cust_id.
If 1, encode existing customer.
If 2, calls encode_cust to create a new customer entry."""

def select_client(cust_table, offer_id):
    
    clear_screen()

    print(tabulate([['Client - offer ID ' + offer_id]], tablefmt="simple_grid"))

    i = True
    while i:
        choice = input('\n1. Existing Customer 2.Encode new customer 3.Search Customer 4.Abort ')
        if choice not in ['1', '2', '3', '4']:
            print('\nEncode the right answer.')
            
        else:
            if choice == '1':
                cust_id = input('\nSélectionner le Customer ID: ')
                
                if cust_id in list(cust_table.keys()):
                    print('Result : {0}, {1}, {2}'.format(cust_table[cust_id]['first_name'], cust_table[cust_id]['last_name'], cust_table[cust_id]['company_name']))
                else:
                    print(cust_id + ' is not in customer list')
                    continue            
                i= False
                
            elif choice == '2':
                cust_table, cust_id = encode_cust(cust_table)
                i = False
                
            elif choice == '3':
                find_cust(cust_table)
                
            else:
                choice  = input('Do you confirm that you want to quit the prg? All previously encoded offers will be saved. (Y,N): ')
                if choice == 'Y':
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                elif choice == 'N':
                    continue
                else:
                    print('Encode Y or N')

    return(cust_table, cust_id)



''' query customer list: allows searching by name or company name using find_cust.'''
def find_cust(cust_table):
    print('\nRecherche client existant')
    print('--------------------------')
    s = input('\nCustomer query :')
    l = []
    for i,j  in cust_table.items():
        if (s in (j.get('last_name') or '')) or (s in (j.get('company_name') or '')) or (s in (j.get('first_name') or '')):
            l.append([i, j['first_name'], j['last_name'], j['company_name']])
    if len(l) > 0:
        print('\n', tabulate(l, headers =['ID', 'Nom', 'Prenom', 'Cny']))
    else:
        print('\n no candidate')



'''Creates new prospect
Prompts user for various customer details like name, address, company information, etc.
Creates a new entry in the customer data with these details.'''

def encode_cust(cust_table):
    print('\nEncodage nouveau prospect')
    print('--------------------------')
    id_list = list(cust_table.keys())
    id_list.sort(reverse = True, key = int)
    cust_id = str(int(id_list[0]) + 1)
    cust_table[cust_id] = {}
    cust_table[cust_id]['titre'] = input('Titre: ')
    cust_table[cust_id]['first_name'] = input('First Name: ')
    cust_table[cust_id]['last_name'] = input('Last Name: ')
    cust_table[cust_id]['company_name'] = input('company name: ')
    cust_table[cust_id]['email'] = input('email: ')
    cust_table[cust_id]['phone_number'] = input('phone number: ')
    cust_table[cust_id]['street'] = input('street: ')
    cust_table[cust_id]['number'] = input('Number: ')
    cust_table[cust_id]['city'] = input('city: ')
    cust_table[cust_id]['postal_code'] = input('postal code: ')
    cust_table[cust_id]['country'] = input('country: ')
    cust_table[cust_id]['customer_type'] = input('customer type: ')
    cust_table[cust_id]['account_creation_date'] = dt.today().date().strftime('%d-%m-%y')
    cust_table[cust_id]['VAT'] = input('VAT: ')
    cust_table[cust_id]['Bank account'] = input('Bank account: ')
    return (cust_table, cust_id)


# In[68]:


''' Takes a customer ID and searches the building data for buildings managed by that customer.
Displays a list of buildings or indicates no buildings found.'''

def query_building(build_table, cust_id):
    l = []
    no_build = False
    for i,j  in build_table.items():
        if cust_id == j['manager']:
            l.append([i, j['street'], j['number'], j['postal_code'], j['city']])
    if len(l) > 0:
        print('\nListe des immeubles gérés par cust nbr '+ cust_id)
        print('--------------------------------------------')
        print((tabulate(l, tablefmt="plain")))
        
    else:
        print("\nPas d'immeuble connu")
        no_build = True
    return(no_build)


'''check if building is managed by client'''
def match_build_cust(x, cust_id, build_table, build_id):
    #print(type(build_table[build_id]['manager']), type(cust_id))
    if build_table[build_id]['manager'] != cust_id:
        print('The building is not managed by this client!')
    else:
        x = False
    return(x)


'''
Retrun build_id
Publish list of building of cust ID.
If list not empty: Asks if the building already exists (Y/N) for a specific offer.
If yes, encode building ID and check if it is managed by cust_id
If no, calls encode_build_table to create a new building entry.'''

def select_building(build_table, cust_id, offer_id):
    
    clear_screen()

    print(tabulate([['Building - Offer ID :' + offer_id]], tablefmt="simple_grid"))

    no_build =  query_building(build_table, cust_id)
    

    if not no_build:
        x=True
        while x:
            exist = input('\n1. Existing Building 2.Encode new building 3.Abort ')

            if exist not in ['1', '2', '3']:
                print('\nEncode the right answer.')
            
            else:  
                if exist == '1':                    
                        build_id = input('\nBuilding ID: ')
                        
                        if build_id in list(build_table.keys()):
                            print('Result : {0}, {1}, {2}, {3}'.format(build_table[build_id]['street'], build_table[build_id]['number'], build_table[build_id]['postal_code'], build_table[build_id]['city']))
                            x = match_build_cust(x, cust_id, build_table, build_id)
                        else:
                            print(build_id + ' is not in building list')
                            continue
                        
                elif exist == '2':
                        build_table, build_id = encode_build_table(build_table, cust_id)
                        x = False

                else:
                    choix  = input('Do you confirm that you want to quit the prg? All previously encoded offers will be saved. (Y,N): ')
                    if choix == 'Y':
                        save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                        os.abort()
                    elif choix == 'N':
                        continue
                    else:
                        print('Encode Y or N')
                
    else :
        build_table, build_id = encode_build_table(build_table, cust_id)


    return(build_table, build_id)


'''Asks if the building address should be the same as the customer's address.
Prompts for building details like street, number, postal code, city, and country.
If a different address is chosen, prompts for those details separately.
Creates a new entry in the building data for the selected customer.'''
def encode_build_table(build_table, cust_id):
    print('\nEncodage nouveau building')
    print('--------------------------')

    id_list = list(build_table.keys())
    id_list.sort(reverse = True, key = int)
    build_id = str(int(id_list[0]) + 1)
    build_table[build_id] = {}

    i = True
    while i:
        idem = input('\nEncode customer address? (Y/N): ')
        new_build = {}
        if idem == 'Y':
                new_build['street'] = cust_table[cust_id]['street']
                new_build['number'] = cust_table[cust_id]['number']
                new_build['postal_code'] = cust_table[cust_id]['postal_code']
                new_build['city'] = cust_table[cust_id]['city']
                new_build['country'] = cust_table[cust_id]['country']
                build_table, i = load_new_building(new_build, build_id, build_table, i, cust_id)  

        elif idem == 'N':    
                new_build['street'] = input('Street: ')
                new_build['number'] = input('Number: ')
                new_build['postal_code'] = input('CP: ')
                new_build['city'] = input('City: ')
                new_build['country'] = input('Country: ')
                build_table, i = load_new_building(new_build, build_id, build_table, i, cust_id)
            
        else:
                print('\nEncode only Y or N')
    return(build_table, build_id)


'load newly encoded building data in build_table'''
def load_new_building(new_build, build_id, build_table, i, cust_id):  
    a = True

    try:
        for j in build_table.values():
            if (new_build['street'] == j['street'] ) and (new_build['number'] == j['number'] ) and (new_build['city'] == j['city'] ):
                print('Building already in buidling list\n' + j['street'] + j['number'])
                a= False
    except:
        pass
        
    if a:
        build_table[build_id]['street'] = new_build['street']
        build_table[build_id]['number'] = new_build['number']
        build_table[build_id]['postal_code'] = new_build['postal_code']
        build_table[build_id]['city'] = new_build['city']
        build_table[build_id]['country'] = new_build['country']
        build_table[build_id]['manager'] = cust_id
        build_table[build_id]['account_creation_date'] = dt.today().date().strftime('%d-%m-%y')
        i=False
    
    return (build_table, i)



'''Allows adding existing products from a product list to the BOM.
Takes user input for product ID and calculates the line price considering quantity and discount.'''
def select_prod_item(prod_table, bom, bom_id):

    a = input('\nProd ID: ')
    
    if a in list(prod_table.keys()):
        print('Results :{0} , {1}'.format(prod_table[a]['description'], str(prod_table[a]['unit_selling_price'])))
        bom_id += 1
        bom[str(bom_id)]= {}
        bom[str(bom_id)]['ID'] = a
        bom[str(bom_id)]['description'] = prod_table[a]['description'] 
        bom[str(bom_id)]['unit_selling_price'] = prod_table[a]['unit_selling_price']
        
        while True:
            try:
                bom[str(bom_id)]['qty'] = int(input('Qty: '))
            except ValueError:
                print('Encode number')
            else: 
                break
            
        while True:
            try:
                z = True
                while z:
                    bom[str(bom_id)]['discount'] = int(input('Discount (%): '))
                    if (bom[str(bom_id)]['discount'] > 100) or  (bom[str(bom_id)]['discount'] < 0):
                            print ('Encode in between 0 and 100')
                    else:
                        z =False
            except ValueError:
                print('Encode number')
            else: 
                break
            
        bom[str(bom_id)]['line price'] = bom[str(bom_id)]['qty'] * bom[str(bom_id)]['unit_selling_price'] * (1 - bom[str(bom_id)]['discount']/100)

    else:
        print(a + ' is not in product list')
        
    return(prod_table, bom, bom_id)


''' query product list: allows searching by description '''
def find_prod(prod_table):
    print('\nRecherche produit existant')
    print('--------------------------')
    s = input('\nProduct query :')
    l = []
    for i,j  in prod_table.items():
        if (s.lower() in (j.get('description') or '').lower()):
            l.append([i, j['description'], j['unit_selling_price'], j['unit_cost']])
    if len(l) > 0:
        print('\n', tabulate(l, headers =['Prod ID', 'Description', 'Price', 'Cost']))
    else:
        print('\n no product')



'''Provides options to add existing items, create new products, erase items, exit BOM creation or abort'''
def select_bom(prod_table, bom, offer_id):
    
    clear_screen()

    print(tabulate([['Bill of Materials - Offer ID :' + offer_id]], tablefmt="simple_grid"))
    
    if len(bom) == 0:
        bom_id =1
    else:
        id_list = list(bom.keys())
        id_list.sort(reverse = True, key = int)
        bom_id = int(id_list[0])

    x = True
    while x:
        choice = input('\n1: Add new item of pricelist, 2:Encode new product, 3:Erase item of the BOM, \n4:Search Product, 5: Exit, 6: Abort \n')
        if choice not in ['1', '2', '3', '4', '5', '6']:
            print('\nEncode the right answer.')
        else:
            if choice == '1':
                prod_table, bom, bom_id = select_prod_item(prod_table, bom, bom_id)
                
            elif choice == '2':
                prod_table, bom, bom_id = creation_prod_item(prod_table, bom, bom_id)
                
            elif choice == '3':
                #display bom items table
                l=[]
                for i, j in bom.items():
                    l.append([i, j['description'], j['unit_selling_price'], j['qty'], j['discount']])
                if len(l) > 0:
                    print('Items in the offer Bill of Materials')
                    print('\n', tabulate(l, headers =['Bom ID', 'Description', 'Price', 'Qty', 'Disc' ]))
                else:
                    print('\n no content')

                #query item to erase
                a = input('\nItem to erase: ')
                b = [str(c) for c in list(bom.keys())]
                if a in b:
                    del(bom[a])
                    print('Item is erased.')
                else:
                    print('Item not in the list.')
                    
            elif choice == '4':
                find_prod(prod_table)
                    
            elif choice == '5' and len(bom) == 0:
                print('\nThere must be at least on item in the Bill of Material')
                
            elif choice == '6':
                choix  = input('Do you confirm that you want to quit the prg? All previously encoded offers will be saved. (Y,N): ')
                if choix == 'Y':
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                elif choix == 'N':
                    continue
                else:
                    print('Encode Y or N')
            
            else:
                x = False

                
    return(prod_table, bom)


'''Creates a new product entry with details like description, selling price, and cost price.
Adds the new product to the BOM with specified quantity and discount.'''
def creation_prod_item(prod_table, bom, bom_id):

    id_list = list(prod_table.keys())
    id_list.sort(reverse = True, key = int)
    prod_id = str(int(id_list[0]) + 1)
    prod_table[prod_id] = {}
    prod_table[prod_id]['description'] = input('\nProd or Service Description: ')

    while True:
        try:
            prod_table[prod_id]['unit_selling_price'] = float(input('Selling Price: '))
        except ValueError:
            print('Encode number')
        else: 
            break
            
    while True:
        try:
            prod_table[prod_id]['unit_cost'] = float(input('Cost: '))
        except ValueError:
            print('Encode number')
        else: 
            break

    bom_id+=1
    bom[str(bom_id)]= {}
    bom[str(bom_id)]['ID'] = prod_id
    bom[str(bom_id)]['description'] = prod_table[prod_id]['description']
    bom[str(bom_id)]['unit_selling_price'] = prod_table[prod_id]['unit_selling_price']

    while True:
        try:
            bom[str(bom_id)]['qty'] = int(input('Qty: '))
        except ValueError:
            print('Encode number')
        else: 
            break

    while True:
        try:
            z = True
            while z:
                bom[str(bom_id)]['discount'] = int(input('Discount (%): '))
                if (bom[str(bom_id)]['discount'] > 100) or  (bom[str(bom_id)]['discount'] < 0):
                        print ('Encode in between 0 and 100')
                else:
                    z =False
        except ValueError:
            print('Encode number')
        else: 
            break

    bom[str(bom_id)]['line price'] = bom[str(bom_id)]['qty'] * bom[str(bom_id)]['unit_selling_price'] * (1 - bom[str(bom_id)]['discount']/100)
            
    return(prod_table, bom, bom_id)


''' query sow list: allows searching by description and paragraph content '''
def find_parag(parag_table):
    print('\nRecherche paragraph existant')
    print('----------------------------')
    s = input('Paragraph query :')
    l = []
    for i,j  in parag_table.items():
        if (s.lower() in (j.get('description') or '').lower() or s.lower() in (j.get('Paragraph') or '').lower()):
            l.append([i, j['description'], j['Paragraph'][:40]])
    if len(l) > 0:
        print('\n', tabulate(l, headers =['Parag ID', 'Description', 'Paragraph']))
    else:
        print('\n no content')


''' Allows adding existing SOW paragraphs or creating new ones.
Provides options to add paragraphs, create new paragraphs, erase paragraphs, or exit SOW creation.'''
def select_content(parag_table, content, offer_id):

    clear_screen()

    print(tabulate([['Statement of Work - Offer ID: ' +  offer_id]], tablefmt="simple_grid"))

    
    if len(content) == 0:
        sow_id = 0
    else:
        id_list = list(content.keys())
        id_list.sort(reverse = True, key = int)
        sow_id = int(id_list[0])
    

    x = True
    while x:
        choice = input('\n1: Add new SOW item , 2:create and add a new SOW item, 3:Erase item of the SOW,\n 4:Search SOW item, 5:Exit, 6: Abort\n')
        if choice not in ['1', '2', '3', '4', '5', '6']:
            print('\nEncode the right answer.')
        else: 
            if choice == '1':
                content, sow_id = select_sow(parag_table, content, sow_id)
                
            elif choice == '2':
                parag_table, content, sow_id = creation_sow_item(parag_table, content, sow_id)
                
                
            elif choice == '3':
                #list parag of the letter

                l = []
                for i,j  in content.items():
                        l.append([i, j['description'], j['Paragraph'][:40]])
                if len(l) > 0:
                    print('List of paragraphs of the offer\n--------------------------------')
                    print('\n', tabulate(l, headers =['Content ID', 'Description', 'Paragraph']))
                else:
                    print('\n no content')

                #erase item      
                a = input('\nItem to erase: ')
                b = [str(c) for c in list(content.keys())]
                if a in b and type(list(content.keys())[0]) == str:
                    del(content[a])
                    print('Item erased.')
                elif a in b and type(list(content.keys())[0]) == int:
                    del(content[int(a)])
                    print('Item erased.')
                else:
                    print('Item not in the list.')

            elif choice == '4':
                find_parag(parag_table)
                    
            elif choice == '5' and len(content) == 0:
                print('\nThere must be at least one item in the statement of work')

            elif choice == '6':
                choix  = input('Do you confirm that you want to quit the prg? All previously encoded offers will be saved. (Y,N): ')
                if choix == 'Y':
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                elif choix == 'N':
                    continue
                else:
                    print('Encode Y or N')
                
            else:
                x = False


    return(parag_table, content)


'''Creates a new SOW paragraph entry with a description.'''
def creation_sow_item(parag_table, content, sow_id):

    id_list = list(parag_table.keys())
    id_list.sort(reverse = True, key = int)
    parag_id = str(int(id_list[0]) + 1)
    parag_table[parag_id] = {}
    parag_table[parag_id]['description'] = input('Prod or Service Description: ')
    parag_table[parag_id]['Paragraph'] = input('SOW paragragh: ')
    elt_count = parag_table[parag_id]['Paragraph'].count('{}')
    print('nbre elt :', elt_count)
    parag_table[parag_id]['elt'] = {}
    for i in range(elt_count):
        parag_table[parag_id]['elt'][i] = input('Element title ' + str(i+1) + ': ')

    sow_id += 1

    content[str(sow_id)] = {}
    content[str(sow_id)]['description'] = parag_table[parag_id]['description']
    content[str(sow_id)]['Paragraph'] = parag_table[parag_id]['Paragraph']
    content[str(sow_id)]['elt'] = {}

    for i,j in parag_table[parag_id]['elt'].items():
            content[str(sow_id)]['elt'][i] = input('encode value of {0} : '.format(j))        
    return(parag_table, content, sow_id)


'''Select SOW item of the list to be added to offer'''
def select_sow(parag_table, content, sow_id):
    
    a = input('\nParagraph ID: ')
    
    if a in list(parag_table.keys()):
        print('Result :', parag_table[a]['description'])
        sow_id += 1

        content[str(sow_id)] = {}
        content[str(sow_id)]['description'] = parag_table[a]['description']
        content[str(sow_id)]['Paragraph'] = parag_table[a]['Paragraph']
        content[str(sow_id)]['elt'] = {}

        for i,j in parag_table[a]['elt'].items():
            content[str(sow_id)]['elt'][i] = input('encode value of {0}: '.format(j))
    else:
        print(a + ' is not in SOW list')

    return(content, sow_id)


# In[80]:


# create offer number
def create_offer_nbr(offers_table):  
    a = list(offers_table.keys())
    b = [i[:8] for i in a]
    c = b.count(dt.today().date().strftime('%d-%m-%y')) + 1
    return (dt.today().strftime('%d-%m-%y') + '-' + str(c))


''' create customer address string for display and letter'''
def create_cl_address(cust_table, cust_id):
    cust_val = (cust_table[cust_id])
    for i in cust_val:
        if cust_val[i] == None:
            cust_val[i] = ' '
    cust_ad = '\t\t\t\t\t\t' + cust_val['company_name'] + '\n\t\t\t\t\t\t' + cust_val['first_name'] + ' ' + cust_val['last_name'] + '\n\t\t\t\t\t\t' + cust_val['street'] + ' ' + cust_val['number']+ '\n\t\t\t\t\t\t' + cust_val['postal_code'] + '' + cust_val['city'] + '\n\t\t\t\t\t\t' + cust_val['country']
    return (cust_ad)


#write offer to docx document
def create_letter(cust_table, cust_id, bom, content, build_table, build_id, offer_id, bom_table, parag_table, path_schema):

    # Create a new Document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier'
    font.size = Pt(10)
    
    # Add logo in the left cell
    logo = doc.add_paragraph()
    logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pict_logo = logo.add_run()
    pict_logo.add_picture('Kanalis-logo-ok.png', width=Inches(1))
    
    # Add sender's address in the right cell
    sender_paragraph = doc.add_paragraph("Kanalys SPRL\nRue de la Croix 41 \n1050 Ixelles", style ='Body Text')
    sender_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Add a line break
    doc.add_paragraph()
    
    # Add recipient's address on the right below
    cust_ad= create_cl_address(cust_table, cust_id) 
    recipient_paragraph = doc.add_paragraph(cust_ad)
    recipient_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    #recipient_address = recipient_paragraph.add_run(cust_ad)
    
    # Add today's date on the right below the recipient's address
    doc.add_paragraph()  # Add some space
    date_paragraph = doc.add_paragraph('Brussels, ' + dt.today().strftime('%B %d, %Y'))
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Add reference section below the date
    doc.add_paragraph()  # Add some space
    reference_paragraph = doc.add_paragraph()
    reference_text = reference_paragraph.add_run("Reference: " + offer_id)
    
    # Add greeting and the main letter content
    doc.add_paragraph()  # Add some space
    greeting_paragraph = doc.add_paragraph("Dear " + cust_table[cust_id]['titre'] + ' ' + cust_table[cust_id]['first_name'] + ' ' + cust_table[cust_id]['last_name'] + ',')
    greeting_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Main text of the letter
    doc.add_paragraph(
        parag_table['10']['Paragraph'].format(build_table[build_id]['street'], build_table[build_id]['number'], build_table[build_id]['postal_code'], build_table[build_id]['city']))


    # Add work description
    # Add a title or heading
    doc.add_heading('Description des travaux', level=1)

    ligne = 30
    
    for j in content.values():
            ligne += round(len(j['Paragraph']))
            paragraph = doc.add_paragraph(j['Paragraph'].format(*[i for i in list(j['elt'].values())]))
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_together

    #Add picture
    try:
        ligne += 17
        doc.add_heading('Schema des travaux', level=1)
        doc.add_picture(path_schema, height=Inches(2))
    except UnrecognizedImageError:
        print('Image non compatible.')
    
    # Add a table with two columns (Index, Paragraph)
    doc.add_heading('Bill of Materials', level=1)
        
    # Add Bom table
    paragraph = doc.add_paragraph(bom_table)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.keep_together
        
    # Add closing remarks
    closing_paragraph = doc.add_paragraph()
    closing_text = closing_paragraph.add_run('\n' + parag_table['11']['Paragraph'] + "\nSincerely,\n\nYour Name\nYour Job Title")
    #closing_text.font.size = Pt(12)
    
    # Save the document
    name_offer = offer_id + '.docx'
    os.chdir('Letters')
    doc.save(name_offer)
    offer_path = 'Letters/' + name_offer
    print("\nLetter template created successfully.")
    os.chdir('..')
    return(offer_path, path_schema)



''' display offer content on screen'''
def display_offre(bom_table, content, offer_id, cust_table, cust_id, parag_table, build_table, build_id):

    clear_screen()

    print(tabulate([['Offer content']], tablefmt="simple_grid"))

    cust_ad= create_cl_address(cust_table, cust_id)
    print(cust_ad + '\n')
    print("Reference: " + offer_id + '\n')
    print('\t\t\t\t\t\t' + dt.today().strftime('%B %d, %Y') + '\n')
    print("Dear " + cust_table[cust_id]['titre'] + ' ' + cust_table[cust_id]['first_name'] + ' ' + cust_table[cust_id]['last_name'] + ',\n')
    print(parag_table['10']['Paragraph'].format(build_table[build_id]['street'], build_table[build_id]['number'], build_table[build_id]['postal_code'], build_table[build_id]['city']) + '\n')
    print('Description des travaux\n------------------\n')
    for j in content.values():
        print(j['Paragraph'].format(*[i for i in list(j['elt'].values())]) + '\n')
    print('Bill of Materials\n------------------\n')
    print (bom_table + '\n')
    print(parag_table['11']['Paragraph'] + '\n')
    print("Sincerely,\n\nYour Name\nYour Job Title")

    print('Size table : ', len(bom_table))




'''create BOM table with pricing that will be integrated in the offer discplay and letter'''
def create_bom_table(bom):
    header = list(list(bom.values())[0].keys())[1:]
    header[1] = 'Unit Price'
    bom_list= []
    for i in bom.values():
        a = list(i.values())[1:]
        bom_list.append(a)
    bom_list
    
    for i in bom_list:
        i[1] = "{:.2f} €".format(i[1])
        i[3] = "{:.0f} %".format(i[3])
        i[4] = "{:.2f} €".format(i[4])
    
    # calcul total
    tot_htva = 0
    tva = 0
    for j in bom.values():
        tot_htva +=  j['line price']
    tva += tot_htva * 0.21 
    tot_tvac = tot_htva * 1.21
    bom_list.append(['-----', '-----', '---', '-----', '-----'])
    bom_list.append([' ', ' ', ' ', 'Total HTVA', "{:.2f} €".format(tot_htva)])
    bom_list.append([' ', ' ', ' ', 'TVA', "{:.2f} €".format(tva)])
    bom_list.append([' ', ' ', ' ', 'Total TVAC', "{:.2f} €".format(tot_tvac)])
    bom_table = tabulate(bom_list, headers = header, tablefmt="orgtbl", colalign = ('left', 'right', 'left' , 'left' ,'right'), floatfmt="{:.2f} €", maxcolwidths=[15, 10, None, None, None])          
    
    return(bom_table)


'''update offer table with new offer of updated offer'''
def create_offer_table(offers_table, offer_id, cust_id, build_id, bom, content, offer_path, path_schema):
    offers_table[offer_id] = {'Customer': cust_id, 'Building': build_id, 'Bom': bom, 'SOW': content,
                              'Path': offer_path, 'Status': 'in approval',
                              'Created': dt.today().date().strftime('%d-%m-%y'),
                              'Schema': path_schema}
    return(offers_table)



#Save updated tables to json
def save_json(cust_table, build_table, offers_table, parag_table, prod_table):
    # Saving the list as a JSON file
    
    with open('json/customers.json', 'w') as file:
        json.dump(cust_table, file, indent=4)
    
    with open('json/offers.json', 'w') as file:
        json.dump(offers_table, file, indent=4)
        
    with open('json/buildings.json', 'w') as file:
        json.dump(build_table, file, indent=4)

    with open('json/sewer_products_services.json', 'w') as file:
        json.dump(prod_table, file, indent=4)

    with open('json/paragraphs.json', 'w') as file:
        json.dump(parag_table, file, indent=4)

    print('Fichiers sauvés')

''' function to change offer content (customer, building, bom items, sow content'''
def change_offer(cust_id, cust_table, build_id, build_table, bom, content, prod_table, parag_table, offer_id, path_schema):

    x = True
    while x:    
        choice = input('\nDo you want to change the content of this offer?\n1(Customer), 2(Building), 3(prod & service), 4(job description),\n 5(schema), 6(display update), 7(create letter and exit)\n')

        if choice not in ['1', '2', '3', '4', '5', '6', '7']:
            print('Encode the right answer.')
        else:
          
            if choice == '1':
                print('\nYour choice: {0}, {1}, {2}, {3}'.format(cust_id, cust_table[cust_id]['company_name'], cust_table[cust_id]['first_name'], cust_table[cust_id]['last_name']))
                cust_table, cust_id = select_client(cust_table, offer_id)
    
                y = True
                while y:
                    y = match_build_cust(y, cust_id, build_table, build_id)
                    if y:
                        no_build = query_building(build_table, cust_id)
                        build_table, build_id = select_building(build_table, cust_id, offer_id)
                        y = False
    
            elif choice == '2':
                print('\nYour choice: {0}, {1}, {2}, {3}'.format(build_id, build_table[build_id]['street'], build_table[build_id]['number'], build_table[build_id]['city']))
                no_build = query_building(build_table, cust_id)
                build_table, build_id = select_building(build_table, cust_id, offer_id)
    
            elif choice == '3':
                header = list(list(bom.values())[0].keys())
                header[1] = 'Unit Price'
                bom_list= []
                for i, j in bom.items():
                    a = list(str(i)) + list(j.values())[1:4]
                    bom_list.append(a)
    
                print('\nYour choice: \n', tabulate(bom_list, headers = header))
                prod_table, bom = select_bom(prod_table, bom, offer_id)
                
            elif choice == '4':
                a = [[i, j['description']] for i,j in content.items()]
                print('\nYour choice: \n', tabulate(a, headers = ['SOW ID', 'Description']))
                parag_table, content = select_content(parag_table, content, offer_id)

            elif choice == '5':
                path_schema = crop_image(offer_id)
                image = Image.open(path_schema)
                image.show()
                
            elif choice == '6':
                bom_table = create_bom_table(bom)
                display_offre(bom_table, content, offer_id, cust_table, cust_id, parag_table, build_table, build_id)
    
            elif choice == '7':       
                x = False
                z = False


    return(z, cust_id, cust_table, build_id, build_table, bom, content, prod_table, parag_table, path_schema)     



''' list of offers'''
def create_list_offers(offers_table, cust_table, status = 'All', s= 'All', date_beg=dt.strptime('01-01-20', '%d-%m-%y'),
                       date_end=dt.today()):

    cu = []
    for i,j  in cust_table.items():
        if (s.lower() in (j.get('last_name') or '').lower())  or (s.lower()  in (j.get('company_name') or '').lower())  or (s.lower()  in (j.get('first_name') or '').lower()):
            cu.append(i)
    t = []
    offer_list = []
    if len(cu) > 0 or s == 'All':
          
        for i, j in offers_table.items():
            if ((status == 'All'or j['Status'] == status )
                and (s == 'All' or j['Customer'] in cu)
                and (dt.strptime(j['Created'], '%d-%m-%y') >= date_beg)
                and (dt.strptime(j['Created'], '%d-%m-%y') <= date_end)):
                                    
                                                                                  
                l =[i]
                offer_list.append(i)
                if cust_table[j['Customer']]['company_name'] == ' ':
                    l.append(cust_table[j['Customer']]['last_name'])
                else:
                    l.append(cust_table[j['Customer']]['company_name'])

                c = 0    
                for a,b in j['Bom'].items():
                    c +=  b['line price']
                l.append(c)
                l.append(j['Status'])
                l.append(j['Created'])
                t.append(l)
    else:
        print('No identified customers.')
        
    print('\nOffer list :')    
    print(tabulate(t, headers = ['Id', 'Client', 'Value', 'Status', 'last update']))
    return(offer_list)



'offer modification process'
def modify_offer(offer_id, cust_table, build_table, prod_table, offers_table, parag_table):
    bom = offers_table[offer_id]['Bom']
    content = offers_table[offer_id]['SOW']
    cust_id = offers_table[offer_id]['Customer']
    build_id = offers_table[offer_id]['Building']
    path_schema = offers_table[offer_id]['Schema']
    
    bom_table = create_bom_table(bom)
    display_offre(bom_table, content, offer_id, cust_table, cust_id, parag_table, build_table, build_id)
    image = Image.open(path_schema)
    image.show()
    z, cust_id, cust_table, build_id, build_table, bom, content, prod_table, parag_table, path_schema = change_offer(cust_id, cust_table, build_id, build_table, bom, content, prod_table, parag_table, offer_id, path_schema)
    bom_table = create_bom_table(bom)
    offer_path, path_schema = create_letter(cust_table, cust_id, bom, content, build_table, build_id, offer_id, bom_table, parag_table, path_schema)
    offers_table = create_offer_table(offers_table, offer_id, cust_id, build_id, bom, content, offer_path, path_schema)
    return (cust_table, build_table, prod_table, offers_table, parag_table)   


'offer creation process'
def create_offer(cust_table, build_table, prod_table, offers_table, parag_table):
    offer_id = create_offer_nbr(offers_table)
    cust_table, cust_id = select_client(cust_table, offer_id)
    build_table, build_id = select_building(build_table, cust_id, offer_id)
    bom = {}
    prod_table, bom = select_bom(prod_table, bom, offer_id)
    content = {}
    parag_table, content = select_content(parag_table, content, offer_id)
    path_schema = crop_image(offer_id)
    z=True
    while z:
        
        bom_table = create_bom_table(bom)
        image = Image.open(path_schema)
        image.show()
        display_offre(bom_table, content, offer_id, cust_table, cust_id, parag_table, build_table, build_id)
        z, cust_id, cust_table, build_id, build_table, bom, content, prod_table, parag_table, path_schema = change_offer(cust_id, cust_table, build_id, build_table, bom, content, prod_table, parag_table, offer_id, path_schema)

    bom_table = create_bom_table(bom)
    offer_path, path_schema = create_letter(cust_table, cust_id, bom, content, build_table, build_id, offer_id, bom_table, parag_table, path_schema)
    offers_table = create_offer_table(offers_table, offer_id, cust_id, build_id, bom, content, offer_path, path_schema)
    return(cust_table, build_table, prod_table, offers_table, parag_table)


'approve/modify offers'
def approve_offer(cust_table, build_table, prod_table, offers_table, parag_table):
    clear_screen()
    print(tabulate([['Offer Approval: ']], tablefmt="simple_grid"))
    create_list_offers(offers_table, cust_table, status = 'in approval')


    x = True
    while x:
        choice = input('\n1: Display Offer, 2: Modify offer, 3:Erase offer ,\n 4:Approve offer, 5:Exit, 6: Abort\n')
        if choice not in ['1', '2', '3', '4', '5', '6']:
            print('\nEncode the right answer.')
        else: 
            if choice == '1':
                offer_id = input('\nOffer ID: ')
                if offer_id not in list(offers_table.keys()):
                    print('Offer does not exist')
                else:
                    bom = offers_table[offer_id]['Bom']
                    content = offers_table[offer_id]['SOW']
                    cust_id = offers_table[offer_id]['Customer']
                    build_id = offers_table[offer_id]['Building']
                    path_schema = offers_table[offer_id]['Schema']
                    bom_table = create_bom_table(bom)
                    display_offre(bom_table, content, offer_id, cust_table, cust_id, parag_table, build_table, build_id)
                    image = Image.open(path_schema)
                    image.show()
                
            elif choice == '2':
                offer_id = input('\nOffer ID: ')
                if offer_id not in list(offers_table.keys()):
                    print('Offer does not exist')
                elif offers_table[offer_id]['Status'] != 'in approval':
                    new_offer_id = create_offer_nbr(offers_table)
                    print('This offer has already been sent. An new one will be created with ID nbr ' + new_offer_id)
                    offers_table['new_offer_id'] = offers_table['new_offer_id']
                    modify_offer(new_offer_id, cust_table, build_table, prod_table, offers_table, parag_table)
                    
                else:
                    modify_offer(offer_id, cust_table, build_table, prod_table, offers_table, parag_table)
            
                
            elif choice == '3':
                offer_id = input('\nOffer to erase: ')
                if offer_id not in list(offers_table.keys()):
                    print('Offer does not exist')
                elif offers_table[offer_id]['Status'] == 'in approval':
                    del(offers_table[offer_id])
                    print('Item erased.')
                elif offers_table[offer_id]['Status'] != 'in approval':
                    print('Already ben sent. May not be deleted')
                else:
                    print("This offer doesn't exist.")

            elif choice == '4':
                offer_id = input('\nOffer to approve: ')
                if offer_id not in list(offers_table.keys()):
                    print('Offer does not exist')
                elif offers_table[offer_id]['Status'] == 'in approval':
                    offers_table[offer_id]['Status'] = 'approved'


                    #collect data for mail
                    cust_id = offers_table[offer_id]['Customer']
                    email = cust_table[cust_id]['email']
                    body = cust_table[cust_id]['titre'] + ' ' + cust_table[cust_id]['first_name'] + ' ' + cust_table[cust_id]['last_name']+ '\n' + parag_table['17']['Paragraph'].format(cust_table[cust_id]['street'], cust_table[cust_id]['number'], cust_table[cust_id]['postal_code'],cust_table[cust_id]['city'])
                    print(offers_table[offer_id]['Path'], type(offers_table[offer_id]['Path']))
                    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', offers_table[offer_id]['Path'], '--outdir', 'Letters'])
                    attach = offers_table[offer_id]['Path'][:-4] + 'pdf'
                    print('attachement', attach)
                    titre = 'Kanalis offer ' + offer_id
                    send_offer(titre, email, body, attach)
                    print('Offer '+ offer_id + ' is sent!')
                else:
                    print('Offer not in approval status')
                    
            elif choice == '6':
                choix  = input('Do you confirm that you want to quit the prg? All previously encoded offers will be saved. (Y,N): ')
                if choix == 'Y':
                    save_json(cust_table, build_table, offers_table, parag_table, prod_table)
                    os.abort()
                elif choix == 'N':
                    continue
                else:
                    print('Encode Y or N')
                
            else:
                x = False


    return(cust_table, build_table, offers_table, parag_table, prod_table)   

'''crop image from visit document'''
def crop_image(offer_id):
    
    clear_screen()
    print(tabulate([['Add Schema- Offer ID: ' +  offer_id]], tablefmt="simple_grid"))
    
    #identify available reports and schema
    print('\nlist of availabe reports and snapshots\n---------------------------------------')
    os.chdir('Chantiers')
    print(os.listdir())

    y = True
    while y:
        path_schema = input('\nEncode file to load: ')
        if path_schema not in os.listdir():
            print('This file does not exist')
        else:
            y = False
    
    type_image = input('\n1.image (jpg, png) , 2.rapport standard (pdf) ')

    x = True
    while x: 
        if type_image not in ['1', '2']:
            print('Encode properly')
        else:

            #store already croped image
            if type_image == '1':
                if path_schema[-3:] not in ['jpg', 'png']:
                    print(' It is not the goof format')
                else:
                    x = False

            #crop image from pdf report
            elif type_image == '2':
                if path_schema[-3:] != 'pdf':
                    print(' It is not the goof format')
                    
                else:
                    reader = PdfReader(path_schema)
                    page = reader.pages[0]
                    #print(page.cropbox.upper_right)

                    writer = PdfWriter()

                    page.cropbox.upper_left = ( 100, 220)
                    page.cropbox.lower_right = ( 350, 50)
                    writer.add_page(page) 
              
                    with open('result.pdf','wb') as fp:
                        writer.write(fp)


                    image = convert_from_path('result.pdf', use_cropbox=True)
                    image[0].save('Image.jpg', 'JPEG')
                    path_schema = 'Image.jpg'
                    os.remove("result.pdf")
                    x= False

    #load final schema in dedicated schma folder and return path
    new_path_schema = 'Schema/schema '+ offer_id + path_schema[-4:]
    os.chdir('..')
    os.rename('Chantiers/{0}'.format(path_schema), new_path_schema)
    return(new_path_schema)
            

def send_offer(titre, email, body, attach):   
    # creates SMTP session
    s = smtplib.SMTP('smtp.gmail.com', 587)
    # start TLS for security
    s.starttls()
    # Authentication
    s.login("beyerman@gmail.com", "invf fzrx lgwn pncd")

    # Create a MIMEText object to represent the email
    msg = MIMEMultipart()
    msg['From'] = 'beyerman@gmail.com'
    msg['To'] = email
    msg['Subject'] = titre
    msg.attach(MIMEText(body, 'plain'))

    with open(attach,'rb') as file:
        # Attach the file with filename to the email
        msg.attach(MIMEApplication(file.read(), Name= titre + '.pdf'))
        
    # sending the mail
    s.sendmail("beyerman@gmail.com", email, msg.as_string())
    # terminating the session
    s.quit()


start(cust_table, build_table, prod_table, offers_table, parag_table)


